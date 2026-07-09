from __future__ import annotations

import html
import re
from dataclasses import dataclass
from pathlib import Path

from app.services.exceptions import KnowledgeParseError, KnowledgeValidationError


SUPPORTED_DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".html", ".htm", ".md", ".txt"}


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    parser: str
    content_type: str
    title: str
    word_count: int


class DocumentParserService:
    def parse_path(self, path: Path) -> ParsedDocument:
        extension = path.suffix.lower()
        if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
            raise KnowledgeValidationError(f"Unsupported document type: {extension}")

        if extension == ".pdf":
            text = self._parse_pdf(path)
            parser = "pypdf"
            content_type = "application/pdf"
        elif extension == ".docx":
            text = self._parse_docx(path)
            parser = "python-docx"
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif extension in {".html", ".htm"}:
            text = self._parse_html(path.read_text(encoding="utf-8", errors="ignore"))
            parser = "html"
            content_type = "text/html"
        elif extension == ".md":
            text = self._parse_markdown(path.read_text(encoding="utf-8", errors="ignore"))
            parser = "markdown"
            content_type = "text/markdown"
        else:
            text = path.read_text(encoding="utf-8", errors="ignore")
            parser = "text"
            content_type = "text/plain"

        cleaned = self._clean_text(text)
        if not cleaned:
            raise KnowledgeParseError("Document contains no extractable text")

        return ParsedDocument(
            text=cleaned,
            parser=parser,
            content_type=content_type,
            title=path.stem,
            word_count=len(cleaned.split()),
        )

    def _parse_pdf(self, path: Path) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise KnowledgeParseError("PDF parser dependency pypdf is not installed") from exc

        try:
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as exc:
            raise KnowledgeParseError("Unable to extract text from PDF") from exc

    def _parse_docx(self, path: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise KnowledgeParseError("DOCX parser dependency python-docx is not installed") from exc

        try:
            document = Document(str(path))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        except Exception as exc:
            raise KnowledgeParseError("Unable to extract text from DOCX") from exc

    def _parse_html(self, raw_html: str) -> str:
        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(raw_html, "html.parser")
            for element in soup(["script", "style", "noscript"]):
                element.decompose()
            return soup.get_text("\n")
        except ImportError:
            text = re.sub(r"(?is)<(script|style).*?>.*?</\\1>", " ", raw_html)
            text = re.sub(r"<[^>]+>", " ", text)
            return html.unescape(text)

    def _parse_markdown(self, text: str) -> str:
        text = re.sub(r"```.*?```", " ", text, flags=re.S)
        text = re.sub(r"`([^`]+)`", r"\1", text)
        text = re.sub(r"!\[[^\]]*\]\([^)]+\)", " ", text)
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
        text = re.sub(r"^#{1,6}\s*", "", text, flags=re.M)
        return text

    def _clean_text(self, text: str) -> str:
        lines = [" ".join(line.split()) for line in text.splitlines()]
        lines = [line for line in lines if line]
        return "\n".join(lines)
